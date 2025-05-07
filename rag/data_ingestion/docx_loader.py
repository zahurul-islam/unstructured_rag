"""
DOCX document loader.
"""

import os
import logging
from typing import Dict, Tuple, Any

# Configure logging
logger = logging.getLogger(__name__)


def load_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Load a DOCX document and extract its text and metadata.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple containing:
            - Extracted text content (str)
            - Document metadata (Dict[str, Any])
    """
    try:
        # Try to use langchain's document loader
        try:
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            
            # Load the DOCX document
            loader = UnstructuredWordDocumentLoader(file_path)
            documents = loader.load()
            
            # Combine text from all pages
            text = " ".join([doc.page_content for doc in documents])
            
            # Create DOCX-specific metadata
            metadata = {
                "source_type": "docx",
                "extraction_method": "unstructured",
                "section_count": len(documents)
            }
            
            return text, metadata
        
        except ImportError:
            logger.warning("UnstructuredWordDocumentLoader not available. Falling back to docx2txt.")
            return load_docx_with_docx2txt(file_path)
    
    except Exception as e:
        logger.error(f"Error loading DOCX with UnstructuredWordDocumentLoader: {str(e)}")
        logger.info("Falling back to docx2txt")
        return load_docx_with_docx2txt(file_path)


def load_docx_with_docx2txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Fallback method to load a DOCX using docx2txt.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple containing:
            - Extracted text content (str)
            - Document metadata (Dict[str, Any])
    """
    try:
        import docx2txt
        
        # Extract text from the DOCX
        text = docx2txt.process(file_path)
        
        # Count paragraphs (approximation)
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        # Create DOCX-specific metadata
        metadata = {
            "source_type": "docx",
            "extraction_method": "docx2txt",
            "paragraph_count": len(paragraphs),
            "word_count": len(text.split())
        }
        
        return text, metadata
    
    except ImportError:
        logger.error("docx2txt not available. Trying python-docx.")
        return load_docx_with_python_docx(file_path)
    
    except Exception as e:
        logger.error(f"Error loading DOCX with docx2txt: {str(e)}")
        logger.info("Trying python-docx")
        return load_docx_with_python_docx(file_path)


def load_docx_with_python_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Fallback method to load a DOCX using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple containing:
            - Extracted text content (str)
            - Document metadata (Dict[str, Any])
    """
    try:
        import docx
        
        # Open the DOCX
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs]
        text = '\n'.join(paragraphs)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        
        # Create DOCX-specific metadata
        metadata = {
            "source_type": "docx",
            "extraction_method": "python-docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "word_count": len(text.split())
        }
        
        # Try to extract document properties
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
        except:
            pass
        
        return text, metadata
    
    except Exception as e:
        logger.error(f"Error loading DOCX with python-docx: {str(e)}")
        
        # Return empty text with error metadata as last resort
        return "", {
            "source_type": "docx",
            "extraction_error": str(e),
            "extraction_success": False
        }
